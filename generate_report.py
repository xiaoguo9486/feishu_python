from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import os
import sys
from datetime import datetime
import re
import glob
from docx.shared import RGBColor
import zipfile
import shutil

# ==============================================
# 配置区
# ==============================================
TEMPLATE_HIGH = "光伏巡检报告模板-高压.docx"
TEMPLATE_LOW  = "光伏巡检报告模板-低压.docx"

PHOTO_FOLDER = "D:/exercise/python/photos/"       # 照片及 Excel 存放目录
# OUTPUT_FOLDER = "../output/生成的报告/"       # 本机测试环境的路径
OUTPUT_FOLDER = "E:/光伏运维/01流水资料/自动生成的报告/"       # 本机测试环境的路径
LOG_DIR = "D:/exercise/python/logs"    # 日志配置（新增）

# 自动匹配 Excel 文件
# # 历史：在当前目录（D:\exercise\python\feishu_pthoto_edit）搜索excel文件
# EXCEL_PATTERN = "*巡检记录-照片xg*.xlsx"
# excel_files = glob.glob(EXCEL_PATTERN)
# 新：在 D:/exercise/python/photos/ 里搜索
# EXCEL_PATTERN = "*巡检记录-照片xg*.xlsx"
EXCEL_PATTERN = "*xg简易自动巡检系统*.xlsx"
EXCEL_SHEET_NAME = "xg简易自动巡检系统"

# #只匹配一个 Excel 和一个对应的 ZIP的代码
# excel_files = glob.glob(os.path.join(PHOTO_FOLDER, EXCEL_PATTERN))
# if not excel_files:
#     raise FileNotFoundError(f"未找到匹配的 Excel 文件，模式：{EXCEL_PATTERN}")
# EXCEL_PATH = excel_files[0]
# print(f"📂 自动匹配 Excel 文件：{EXCEL_PATH}")

# 图片解压设置
ZIP_PATTERN = "*巡检*附件*.zip"                   # 飞书下载的压缩包匹配规则
IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp'}

# ==============================================
# 邮件通知配置（请根据实际环境修改）
# ==============================================
SMTP_SERVER = "smtp.qq.com"        # 公司邮件服务器地址
SMTP_PORT = 587                         # 通常 465(SSL) 或 587(TLS)
SMTP_USER = "2604182970@qq.com"  # 发件邮箱
SMTP_PASSWORD = "tjeqnwmhwyiyecbg"         # 邮箱密码或授权码
MAIL_TO = "330951244@qq.com,xiaoguo9486@163.com"           # 收件人，多个用英文逗号分隔
MAIL_SUBJECT = "光伏巡检报告已生成"      # 邮件主题

# TEMPLATE_PATH = "光伏巡检报告模板-高压.docx"
# EXCEL_PATH = "金田铜业报告-YH光伏巡检-照片巡检.xlsx"
# PHOTO_FOLDER = "D:/exercise/python/photos/"
# OUTPUT_FOLDER = "../output/生成的报告/"
# EXCEL_SHEET_NAME = "xg简易自动巡检系统"

# 图片统一设置
PHOTO_WIDTH = Inches(2.2)  # ✅ 缩小至2.0英寸，保证三张图能在A4表格中并排
PHOTO_PER_ROW = 3
DEBUG_MODE = False

# 搜索所有匹配的 Excel 和 ZIP 文件（用于后续分组）
all_excel_files = glob.glob(os.path.join(PHOTO_FOLDER, EXCEL_PATTERN))
all_zip_files = glob.glob(os.path.join(PHOTO_FOLDER, ZIP_PATTERN))

if not all_excel_files:
    raise FileNotFoundError(f"未找到匹配的 Excel 文件，模式：{os.path.join(PHOTO_FOLDER, EXCEL_PATTERN)}")
print(f"📂 找到 {len(all_excel_files)} 个 Excel 文件，{len(all_zip_files)} 个 ZIP 文件")

# 调试模式：关闭后只输出关键日志
DEBUG_MODE = False

# ==============================================
# 全局映射表（与原来完全相同，省略，请保留原有代码）
# ==============================================
STATUS_MAP = {
    "status_pv_roof": [1, 2, 3, 4, 5],
    "status_pv_module": [6, 7, 8, 9, 10],
    "status_pv_inverter": [11, 12, 13, 14, 15, 16, 17],
    "status_pv_combiner": [18, 19, 20, 21, 22, 23, 24, 25],
    "status_pv_cable": [26, 27, 28],
    "status_pv_fire": [29],
    "status_lv_sign": [31, 32],
    "status_lv_env": [33, 46],
    "status_lv_grid_cabinet": [34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44],
    "status_lv_grid_point": [45],
    "status_lv_comm": [47],
    "status_lv_fire": [48],
    "status_lv_finish": [49],
    "status_hv_box": [50],
    "status_hv_sign": [51],
    "status_hv_transformer": [52, 53, 54, 55, 56, 57],
    "status_hv_station": [58, 59, 60, 61, 62, 63],
    "status_hv_svg": [64],
    "status_hv_control": [65, 66, 67, 68, 69],
    "status_hv_fire": [70],
    "status_hv_finish": [71]
}

ABNORMAL_ITEM_MAP = {
    1: ("光伏区", "屋面情况", "进出入口标识"),
    2: ("光伏区", "屋面情况", "进出入口通道（爬梯）"),
    3: ("光伏区", "屋面情况", "屋面安全防护"),
    4: ("光伏区", "屋面情况", "屋面天沟情况，漏水情况，彩钢瓦情况"),
    5: ("光伏区", "屋面情况", "屋面污染情况"),
    6: ("光伏区", "组件情况", "组件-组件破损/移动/缺失"),
    7: ("光伏区", "组件情况", "组件-手持热成像对热斑及隐裂检查"),
    8: ("光伏区", "组件情况", "组件-遮阴/杂物/积尘检查"),
    9: ("光伏区", "组件情况", "组件-固定件破损缺失松动"),
    10: ("光伏区", "组件情况", "组件-接地工艺"),
    11: ("光伏区", "逆变器情况", "安装固定/支架螺丝是否松动生锈"),
    12: ("光伏区", "逆变器情况", "面板运行显示是否正常"),
    13: ("光伏区", "逆变器情况", "运行日志是否正常"),
    14: ("光伏区", "逆变器情况", "外壳是否存在破损生锈、编号模糊"),
    15: ("光伏区", "逆变器情况", "交流电缆是否虚接、破损，直流MC4接头是否异常"),
    16: ("光伏区", "逆变器情况", "热成像异常高温检查"),
    17: ("光伏区", "逆变器情况", "接地是否正常"),
    18: ("光伏区", "汇流箱情况", "安装固定/支架螺丝"),
    19: ("光伏区", "汇流箱情况", "箱体外壳生锈/编号"),
    20: ("光伏区", "汇流箱情况", "箱体进水/密封条"),
    21: ("光伏区", "汇流箱情况", "保险丝检查"),
    22: ("光伏区", "汇流箱情况", "进出线电流"),
    23: ("光伏区", "汇流箱情况", "线缆接头松动脱落"),
    24: ("光伏区", "汇流箱情况", "热成像异常高温检查"),
    25: ("光伏区", "汇流箱情况", "接地是否正常"),
    26: ("光伏区", "线路情况", "电缆桥架是否松动、破损，盖板是否盖严，接地是否正常"),
    27: ("光伏区", "线路情况", "线缆是否老化、破损、未做保护、进出孔洞封堵是否异常"),
    28: ("光伏区", "线路情况", "水管是否存在异常"),
    29: ("光伏区", "消防设施", "消防设施/灭火器"),
    30: ("光伏区", "巡视收尾", "巡视完成后，检查所有的设备及通道的门、窗、口是否关闭"),
    31: ("低压配电室", "标识情况", "指示牌/警示牌完好"),
    32: ("低压配电室", "标识情况", "一二次图完好"),
    33: ("低压配电室", "环境情况", "杂物/环境情况"),
    34: ("低压配电室", "并网柜情况", "柜体外壳"),
    35: ("低压配电室", "并网柜情况", "仪表指示灯"),
    36: ("低压配电室", "并网柜情况", "门锁"),
    37: ("低压配电室", "并网柜情况", "微机保护"),
    38: ("低压配电室", "并网柜情况", "光伏终端/计量表"),
    39: ("低压配电室", "并网柜情况", "塑壳/框架断路器"),
    40: ("低压配电室", "并网柜情况", "隔离开关合位"),
    41: ("低压配电室", "并网柜情况", "铜牌"),
    42: ("低压配电室", "并网柜情况", "电缆接头"),
    43: ("低压配电室", "并网柜情况", "线缆孔洞封堵"),
    44: ("低压配电室", "并网柜情况", "热成像高温检查"),
    45: ("低压配电室", "并网点情况", "并网点检查"),
    46: ("低压配电室", "环境情况", "空调/温度检查"),
    47: ("低压配电室", "通讯情况", "通讯设备检查"),
    48: ("低压配电室", "消防设施", "消防设施/灭火器"),
    49: ("低压配电室", "巡视收尾", "巡视检查门窗关闭"),
    50: ("高压配电室", "箱体情况", "箱体/门锁/密封条/环境"),
    51: ("高压配电室", "标识情况", "标识牌/警示牌/一二次图"),
    52: ("高压配电室", "升压箱变情况", "高压进出线柜面板指示灯、柜体、温度、照明是否异常"),
    53: ("高压配电室", "升压箱变情况", "变压器是否存在异响，温度、冷却风机是否异常"),
    54: ("高压配电室", "升压箱变情况", "低压进出线柜面板指示灯、柜体、温度、照明是否异常"),
    55: ("高压配电室", "升压箱变情况", "低压进出线柜微机保护、箱变测控装置是否异常"),
    56: ("高压配电室", "升压箱变情况", "UPS、控制变压器是否异常"),
    57: ("高压配电室", "升压箱变情况", "参照低压并网柜检查"),
    58: ("高压配电室", "汇集站情况", "表计操作面板指示灯显示是否异常"),
    59: ("高压配电室", "汇集站情况", "进线保护装置、站用变保护装置是否正常运行"),
    60: ("高压配电室", "汇集站情况", "刀闸位置及闭锁"),
    61: ("高压配电室", "汇集站情况", "带电显示装置开关指示灯"),
    62: ("高压配电室", "汇集站情况", "柜体、温度、照明是否异常，运行声音是否正常"),
    63: ("高压配电室", "汇集站情况", "空调是否异常"),
    64: ("高压配电室", "SVG情况", "外观/门锁/密封条/面板/异响"),
    65: ("高压配电室", "二次控制室情况", "面板指示灯、柜体、温度、照明是否异常"),
    66: ("高压配电室", "二次控制室情况", "储能电池是否异常"),
    67: ("高压配电室", "二次控制室情况", "微机保护有无告警"),
    68: ("高压配电室", "二次控制室情况", "UPS是否存在异常"),
    69: ("高压配电室", "二次控制室情况", "空调是否异常"),
    70: ("高压配电室", "消防设施", "消防设施/灭火器"),
    71: ("高压配电室", "巡视收尾", "巡视检查门窗关闭")
}

PHOTO_GROUP_MAP = {
    "photo_group_1": [1, 2, 3, 4, 5],
    "photo_group_2": [6, 7, 8, 9, 10],
    "photo_group_3": [11, 12, 13, 14, 15, 16, 17],
    "photo_group_4": [18, 19, 20, 21, 22, 23, 24, 25],
    "photo_group_5": [26, 27, 28],
    "photo_group_6": [29],
    "photo_group_7": [30],
    "photo_group_8": [31, 32],
    "photo_group_9": [33, 46],
    "photo_group_10": [34, 35, 36, 37, 38, 39, 40, 41, 42, 43, 44],
    "photo_group_11": [45],
    "photo_group_12": [47],
    "photo_group_13": [48],
    "photo_group_14": [49],
    "photo_group_15": [50],
    "photo_group_16": [51],
    "photo_group_17": [52, 53, 54, 55, 56, 57],
    "photo_group_18": [58, 59, 60, 61, 62, 63],
    "photo_group_19": [64],
    "photo_group_20": [65, 66, 67, 68, 69],
    "photo_group_21": [70],
    "photo_group_22": [71]
}


# ==============================================
# 辅助函数
# ==============================================
def replace_text_placeholders(doc, data):
    """替换所有文本占位符，并处理<br/>换行"""
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                if isinstance(value, str):
                    value = value.replace("<br/>", "\n").replace("<br>", "\n")
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    for table in doc.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                for key, value in data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in cell.text:
                        if isinstance(value, str):
                            value = value.replace("<br/>", "\n").replace("<br>", "\n")
                        cell.text = cell.text.replace(placeholder, str(value))


def generate_status_summary(row):
    """生成状态汇总数据"""
    status_data = {}
    for status_key, item_numbers in STATUS_MAP.items():
        has_abnormal = False
        for item_num in item_numbers:
            status_col = f"item_{item_num}_status"
            if status_col in row and str(row[status_col]).strip() == "异常":
                has_abnormal = True
                break
        status_data[status_key] = "异常" if has_abnormal else "正常"
    return status_data


def parse_abnormal_summary(summary_text):
    """解析异常汇总文本"""
    abnormal_list = []
    if pd.isna(summary_text) or not summary_text.strip():
        return abnormal_list

    items = [item.strip() for item in summary_text.split(";") if item.strip()]
    for item in items:
        if "." in item:
            num_part, desc_part = item.split(".", 1)
            try:
                item_num = int(num_part.strip())
                desc = desc_part.strip()
                if item_num in ABNORMAL_ITEM_MAP:
                    area, category, inspection_item = ABNORMAL_ITEM_MAP[item_num]
                    abnormal_list.append({
                        "area": area,
                        "category": category,
                        "item": inspection_item,
                        "desc": desc
                    })
            except ValueError:
                continue
    return abnormal_list


def fill_abnormal_table(doc, abnormal_list):
    """填充异常明细循环表格"""
    abnormal_table = doc.tables[2]
    template_row = abnormal_table.rows[1]

    if not abnormal_list:
        for cell in template_row.cells:
            cell.text = "无异常"
        return

    first_abnormal = abnormal_list[0]
    template_row.cells[0].text = first_abnormal["area"]
    template_row.cells[1].text = first_abnormal["category"]
    template_row.cells[2].text = first_abnormal["item"]
    template_row.cells[3].text = first_abnormal["desc"]

    for abnormal in abnormal_list[1:]:
        new_row = abnormal_table.add_row()
        for i in range(4):
            new_row.cells[i].paragraphs[0].style = template_row.cells[i].paragraphs[0].style
        new_row.cells[0].text = abnormal["area"]
        new_row.cells[1].text = abnormal["category"]
        new_row.cells[2].text = abnormal["item"]
        new_row.cells[3].text = abnormal["desc"]


def extract_column_number(col_str):
    """从列名开头提取数字，支持 1、6-1、12-3 格式"""
    match = re.match(r'^(\d+)(-\d+)?', col_str.strip())
    if match:
        return int(match.group(1))
    return None


def insert_photo_groups(doc, excel_row, df_columns):
    """插入所有图片分组（✅ 利用反馈列精准定位照片列）"""
    processed_groups = set()
    all_cells_found = []

    # ---------- 扫描模板中的所有图片占位符 ----------
    for table_idx, table in enumerate(doc.tables):
        for row_idx, table_row in enumerate(table.rows):
            for col_idx, cell in enumerate(table_row.cells):
                group_name = None
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        blips = run._r.xpath('.//a:blip')
                        if blips:
                            cNvPr_list = run._r.xpath('.//pic:cNvPr')
                            if cNvPr_list:
                                cNvPr = cNvPr_list[0]  # ✅ 修复变量缺失
                                alt_text = cNvPr.attrib.get("descr", "")
                                if alt_text.startswith("{{图片:"):
                                    group_name = alt_text[5:-2]
                                    break
                    if group_name:
                        break
                if group_name:
                    all_cells_found.append((table_idx, row_idx, col_idx, group_name))

    print(f"\n🔍 [诊断] 扫描到 {len(all_cells_found)} 个图片占位符")

    # ---------- 构建反馈列索引表 ----------
    feedback_index_map = {}  # {巡检项编号: 反馈列在df_columns中的索引}
    for idx, col in enumerate(df_columns):
        col_str = str(col).strip()
        if col_str.startswith("照片问题反馈"):
            try:
                num = int(col_str.replace("照片问题反馈", "").strip())
                feedback_index_map[num] = idx
            except:
                pass

    # ---------- 照片列查找函数 ----------
    def find_photo_col(photo_num):
        """根据巡检项编号查找对应的巡检照片列"""
        # 方法1：通过反馈列定位（反馈列后面紧跟着的就是照片列）
        if photo_num in feedback_index_map:
            fb_idx = feedback_index_map[photo_num]
            if fb_idx + 1 < len(df_columns):
                next_col = df_columns[fb_idx + 1]
                col_str = str(next_col).strip()
                if "巡检照片" in col_str:
                    return next_col
        # 方法2：后备，使用数字匹配（兼容前19项）
        for col in df_columns:
            col_str = str(col).strip()
            col_num = extract_column_number(col_str)
            if col_num == photo_num and "巡检照片" in col_str:
                return col
        return None

    # ---------- 处理每个占位符 ----------
    for table_idx, row_idx, col_idx, group_name in all_cells_found:
        cell = doc.tables[table_idx].rows[row_idx].cells[col_idx]
        print(f"\n📌 [处理] 分组={group_name} (表格{table_idx},行{row_idx},列{col_idx})")

        if group_name in processed_groups:
            print("   ⏭️  已处理过，跳过")
            continue
        if group_name not in PHOTO_GROUP_MAP:
            print(f"   ⚠️  未知分组 {group_name}")
            processed_groups.add(group_name)
            continue

        # 清空单元格
        for paragraph in cell.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        cell.add_paragraph()

        # 收集图片文件名
        photo_filenames = []
        for photo_num in PHOTO_GROUP_MAP[group_name]:
            photo_col = find_photo_col(photo_num)
            if not photo_col:
                print(f"   ❓ 巡检项 {photo_num} 未找到照片列")
                continue
            if pd.isna(excel_row[photo_col]):
                print(f"   ℹ️  巡检项 {photo_num} 对应单元格为空")
                continue
            filenames = str(excel_row[photo_col]).split(",")
            filenames = [fn.strip() for fn in filenames if fn.strip()]
            photo_filenames.extend(filenames)
            if filenames:
                print(f"   📷 巡检项 {photo_num} 收集到 {len(filenames)} 个文件")

        # 插入图片
        if photo_filenames:
            print(f"   ➡️  准备插入 {len(photo_filenames)} 张图片")
            current_paragraph = cell.paragraphs[0]
            current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            success_count = 0
            for i, fname in enumerate(photo_filenames):
                path = os.path.join(PHOTO_FOLDER, fname)
                if not os.path.exists(path):
                    print(f"      ❌ 图片不存在: {path}")
                    continue
                run = current_paragraph.add_run()
                run.add_picture(path, width=PHOTO_WIDTH)
                success_count += 1
                if (i + 1) % PHOTO_PER_ROW == 0 and i != len(photo_filenames) - 1:
                    current_paragraph = cell.add_paragraph()
                    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"   ✅ 成功插入 {success_count} 张")
        else:
            print(f"   ⚠️  无图片文件")

        processed_groups.add(group_name)

    print("\n📊 [诊断] 已处理分组：", sorted(processed_groups))

# ==============================================
# 日志工具（新增：将控制台输出同步写入文件）
# ==============================================
class Tee:
    """同时输出到控制台和日志文件的类"""
    def __init__(self, file_path):
        self.console = sys.stdout
        self.log_file = open(file_path, 'w', encoding='utf-8')

    def write(self, message):
        self.console.write(message)
        self.log_file.write(message)

    def flush(self):
        self.console.flush()
        self.log_file.flush()

    def close(self):
        self.log_file.close()

def setup_logging():
    """初始化日志：创建目录，返回日志文件路径和Tee对象"""
    os.makedirs(LOG_DIR, exist_ok=True)
    log_filename = f"process_{datetime.now().strftime('%Y%m%d%H%M%S')}.log"
    log_path = os.path.join(LOG_DIR, log_filename)
    tee = Tee(log_path)
    sys.stdout = tee
    return tee, log_path


def send_mail(smtp_server, smtp_port, user, password, to_addrs, subject, body):
    import smtplib
    from email.mime.text import MIMEText

    print("[DEBUG] send_mail 被调用")
    if isinstance(to_addrs, str):
        to_addrs = [addr.strip() for addr in to_addrs.split(',') if addr.strip()]
    if not to_addrs:
        print("[DEBUG] 收件人地址为空，返回失败")
        return False, "收件人地址为空"

    try:
        print("[DEBUG] 开始构建 MIMEText")
        msg = MIMEText(body, 'plain', 'utf-8')
        msg['From'] = user
        msg['To'] = ', '.join(to_addrs)
        msg['Subject'] = subject
        print(f"[DEBUG] 邮件构建完成，主题: {subject}, 收件人: {to_addrs}")

        port = int(smtp_port)
        print(f"[DEBUG] 准备连接服务器 {smtp_server}:{port}")
        if port == 465:
            server = smtplib.SMTP_SSL(smtp_server, port, timeout=15)
            print("[DEBUG] 使用 SMTP_SSL 连接")
        else:
            server = smtplib.SMTP(smtp_server, port, timeout=15)
            server.ehlo()
            if port == 587:
                server.starttls()
                server.ehlo()
            print(f"[DEBUG] 使用普通 SMTP 连接 (端口 {port})")

        print("[DEBUG] 准备登录")
        server.login(user, password)
        print("[DEBUG] 登录成功，准备发送")
        server.sendmail(user, to_addrs, msg.as_bytes())
        print("[DEBUG] 发送成功，准备退出")
        server.quit()
        return True, f"邮件已发送至 {', '.join(to_addrs)}"
    except Exception as e:
        import traceback
        print("[DEBUG] 发生异常：")
        traceback.print_exc()
        return False, f"邮件发送失败: {e}"

""" 面对可能存在的文件名是【YH光伏巡检-照片巡检_xg简易自动巡检系统_照片-昨日 (1).xlsx】问题
处理逻辑是：逻辑是：
对于每个 Excel 文件，先去掉它自身可能存在的 (数字) 后缀，得到 clean_base（文件名没有 (1)，所以 clean_base 就等于 YH光伏巡检...照片-昨日）。
正则 r'\s*\(\d+\)(?=\.zip$)' 会去掉 .zip 前面可能存在的 (1)、(2) 等，无论 ZIP 有没有被重命名都能正确匹配
同时处理了 Excel 和 ZIP 双方的重名后缀
"""
def find_zip_for_excel(excel_path):
    """根据 Excel 文件名查找配套 ZIP（适配 Windows 对 Excel 和 ZIP 都可能添加的 (1) 等重名后缀）"""
    import re
    base = os.path.splitext(os.path.basename(excel_path))[0]
    # 去掉 Excel 末尾可能存在的 (数字) 得到 clean_base
    clean_base = re.sub(r'\s*\(\d+\)\s*$', '', base)

    # 遍历所有已知 ZIP 文件，寻找匹配项
    for zip_full in all_zip_files:          # all_zip_files 是全局变量
        zip_name = os.path.basename(zip_full)
        # 去掉 ZIP 文件名末尾可能的 (数字) 后缀（例如 _附件(1).zip → _附件.zip）
        zip_clean = re.sub(r'\s*\(\d+\)(?=\.zip$)', '', zip_name, flags=re.IGNORECASE)
        # 期望的文件名
        expected = clean_base + "_附件.zip"
        if zip_clean == expected:
            return zip_full
    return None

# ==============================================
# 新增：解压图片函数
# ==============================================
def unpack_photos():
    """
    自动解压 PHOTO_FOLDER 下的飞书附件 zip，将所有图片平铺到 PHOTO_FOLDER 根目录。
    如果没有找到压缩包，则跳过。
    """
    search_path = os.path.join(PHOTO_FOLDER, ZIP_PATTERN)
    zip_files = glob.glob(search_path)
    if not zip_files:
        print("ℹ️  未找到需解压的压缩包，跳过解压步骤。")
        return

    print(f"📦 找到 {len(zip_files)} 个压缩包，开始解压...")
    for zip_path in zip_files:
        print(f"   ⏳ 正在处理：{os.path.basename(zip_path)}")
        temp_dir = os.path.join(PHOTO_FOLDER, "_temp_unpack")
        if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        os.makedirs(temp_dir, exist_ok=True)

        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(temp_dir)

        moved_count = 0
        for root, dirs, files in os.walk(temp_dir):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in IMAGE_EXTENSIONS:
                    src = os.path.join(root, file)
                    dst = os.path.join(PHOTO_FOLDER, file)
                    # 处理重名
                    if os.path.exists(dst):
                        base, ext_orig = os.path.splitext(file)
                        counter = 1
                        while True:
                            new_name = f"{base}_{counter}{ext_orig}"
                            dst = os.path.join(PHOTO_FOLDER, new_name)
                            if not os.path.exists(dst):
                                break
                            counter += 1
                        print(f"      ⚠️  重名处理：{file} → {new_name}")
                    shutil.move(src, dst)
                    moved_count += 1

        shutil.rmtree(temp_dir)
        print(f"   ✅ 解压完成，移动了 {moved_count} 张图片")

        # 可选：删除已处理的压缩包（取消下一行注释即可删除）
        # os.remove(zip_path)
        # print(f"   🗑️  已删除压缩包")

    print("✅ 所有压缩包处理完毕。")

def set_abnormal_status_red(doc):
    """
    遍历文档所有表格，将状态为「异常」的单元格文字设置为 加粗+红色
    同时处理「状态汇总表」和「巡检明细表」
    """
    # 标准红色，可根据需要调整RGB数值
    red_color = RGBColor(255, 0, 0)

    for table in doc.tables:
        # 遍历表格所有行（跳过表头行，从第2行开始）
        for row_idx, row in enumerate(table.rows):
            if row_idx == 0:
                continue

            # --------------------------
            # 适配两种表格的状态列位置
            # 状态汇总表：第3列（索引2）
            # 巡检明细表：第3列（索引2）
            # 两者位置一致，统一处理
            # --------------------------
            cell = row.cells[2]
            status_text = cell.text.strip()

            # 只处理内容为「异常」的单元格
            if status_text == "异常":
                # 遍历单元格内所有段落、所有文本块（run）
                # 格式是绑定在run上的，必须逐个设置才能全部生效
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True  # 加粗
                        run.font.color.rgb = red_color  # 标红


# ==============================================
# 主函数
# ==============================================
def main():
    # ---------- 日志初始化 ----------
    tee, log_path = setup_logging()
    print(f"📝 日志文件：{log_path}")

    # ---------- 第 0 步：自动解压所有图片 ----------
    unpack_photos()

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    total_reports = 0
    missing_zip_list = []          # 记录缺少 ZIP 的 Excel
    summary_details = []           # 用于邮件内容

    # ---------- 遍历每一个 Excel 文件 ----------
    for excel_path in all_excel_files:
        print(f"\n{'='*60}")
        print(f"📂 正在处理 Excel：{os.path.basename(excel_path)}")

        # 检查配套 ZIP
        zip_path = find_zip_for_excel(excel_path)
        if not zip_path:
            missing_zip_list.append(os.path.basename(excel_path))
            print("⚠️  未找到配套压缩包，将仅基于已有图片生成报告")

        print("📥 正在读取Excel数据...")
        df = pd.read_excel(excel_path, sheet_name=EXCEL_SHEET_NAME)
        df.columns = df.columns.str.strip()

        excel_reports = 0
        for index, excel_row in df.iterrows():
            print(f"\n🔄 正在生成第 {index+1}/{len(df)} 份报告...")

            grid_type = str(excel_row["并网类型"]).strip()
            if "高压" in grid_type:
                template_path = TEMPLATE_HIGH
            elif "低压" in grid_type:
                template_path = TEMPLATE_LOW
            else:
                print(f"   ⚠️  无法识别并网类型 '{grid_type}'，默认使用高压模板")
                template_path = TEMPLATE_HIGH

            print(f"   📄 使用模板：{os.path.basename(template_path)}")
            doc = Document(template_path)

            # 1. 基础数据（保持不变）
            data = {
                "entry_time": str(excel_row["录入时间"]).split(" ")[0] if not pd.isna(excel_row["录入时间"]) else "",
                "data_time": str(excel_row["数据时间"]).split(" ")[0] if not pd.isna(excel_row["数据时间"]) else "",
                "area": str(excel_row["所属区域"]) if not pd.isna(excel_row["所属区域"]) else "",
                "station_name": str(excel_row["站点名称"]) if not pd.isna(excel_row["站点名称"]) else "",
                "grid_type": str(excel_row["并网类型"]) if not pd.isna(excel_row["并网类型"]) else "",
                "roof_type": str(excel_row["屋面类型"]) if not pd.isna(excel_row["屋面类型"]) else "",
                "inspector": str(excel_row["录入人"]) if not pd.isna(excel_row["录入人"]) else "",
                "generate_date": datetime.now().strftime("%Y/%m/%d"),
                "inspection_result": "正常",
                "problem_summary": str(excel_row["照片问题反馈汇总"]) if not pd.isna(excel_row["照片问题反馈汇总"]) else ""
            }

            # ---------- 构建反馈列索引表（与图片查找共用） ----------
            feedback_index_map = {}
            for idx, col in enumerate(df.columns):
                col_str = str(col).strip()
                if col_str.startswith("照片问题反馈"):
                    try:
                        num = int(col_str.replace("照片问题反馈", "").strip())
                        feedback_index_map[num] = idx
                    except:
                        pass

            # 2. 巡检项数据（保持不变）
            for item_num in range(1, 72):
                status_col = None
                if item_num in feedback_index_map:
                    fb_idx = feedback_index_map[item_num]
                    if fb_idx - 1 >= 0:
                        prev_col = df.columns[fb_idx - 1]
                        prev_str = str(prev_col).strip()
                        if "巡检照片" not in prev_str and "照片问题反馈" not in prev_str:
                            status_col = prev_col
                if not status_col:
                    for col in df.columns:
                        col_str = str(col).strip()
                        col_num = extract_column_number(col_str)
                        if col_num == item_num and "-巡检照片" not in col_str and "照片问题反馈" not in col_str:
                            status_col = col
                            break
                remark_col = f"照片问题反馈{item_num}"
                data[f"item_{item_num}_status"] = str(excel_row[status_col]) if status_col and not pd.isna(
                    excel_row[status_col]) else "正常"
                data[f"item_{item_num}_remark"] = str(excel_row[remark_col]) if remark_col in df.columns and not pd.isna(
                    excel_row[remark_col]) else ""

            # 3. 状态汇总
            status_data = generate_status_summary(data)
            data.update(status_data)

            # 4. 替换文本
            replace_text_placeholders(doc, data)

            # 5. 异常明细
            abnormal_list = parse_abnormal_summary(data["problem_summary"])
            fill_abnormal_table(doc, abnormal_list)

            # 6. 插入图片
            insert_photo_groups(doc, excel_row, df.columns)

            # 异常状态标红加粗
            set_abnormal_status_red(doc)

            # 7. 保存报告
            station_name = data["station_name"].replace("/", "-").replace("\\", "-")
            data_date = data["data_time"].replace("/", "-")
            output_filename = f"{station_name}_{data_date}_{grid_type}巡检报告.docx"
            output_path = os.path.join(OUTPUT_FOLDER, output_filename)
            doc.save(output_path)
            print(f"✅ 已生成：{output_path}")
            excel_reports += 1

        print(f"✅ 本 Excel 生成 {excel_reports} 份报告")
        total_reports += excel_reports
        summary_details.append((os.path.basename(excel_path), excel_reports, bool(zip_path)))

    # ========== 汇总输出 ==========
    print(f"\n{'='*60}")
    print(f"🎉 全部处理完成！共处理 {len(all_excel_files)} 个 Excel 文件，生成 {total_reports} 份报告")
    print(f"📂 报告保存路径：{os.path.abspath(OUTPUT_FOLDER)}")
    if missing_zip_list:
        print(f"⚠️  以下 Excel 文件缺少配套压缩包：")
        for name in missing_zip_list:
            print(f"   - {name}")

    # ========== 发送邮件通知 ==========
    mail_body = f"本次巡检报告已生成完毕。\n"
    mail_body += f"处理文件数：{len(all_excel_files)}\n"
    mail_body += f"生成报告总数：{total_reports}\n"
    mail_body += f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    mail_body += f"报告存放路径：{os.path.abspath(OUTPUT_FOLDER)}\n"
    for excel_name, count, has_zip in summary_details:
        zip_status = "有压缩包" if has_zip else "无压缩包（使用已有图片）"
        mail_body += f"  · {excel_name}：生成 {count} 份报告，{zip_status}\n"
    if missing_zip_list:
        mail_body += f"\n注意：以下 Excel 缺少配套 ZIP：\n"
        for name in missing_zip_list:
            mail_body += f"  - {name}\n"

    success, msg = send_mail(SMTP_SERVER, SMTP_PORT, SMTP_USER, SMTP_PASSWORD,
                             MAIL_TO, MAIL_SUBJECT, mail_body)
    if success:
        print(f"📧 邮件通知发送成功：{msg}")
    else:
        print(f"⚠️ 邮件通知发送失败：{msg}")

    # 关闭日志文件
    sys.stdout = tee.console
    tee.log_file.close()

if __name__ == "__main__":
    main()